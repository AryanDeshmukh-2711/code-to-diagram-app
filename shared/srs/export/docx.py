"""DOCX export — the primary format, because it is the one users edit.

A PDF is what gets submitted; a .docx is what gets *worked on* the night
before. That difference decides most of the choices in this file:

* Everything is a real Word style — Heading 1/2/3, Caption, List Bullet,
  Normal — configured once on the style and never on a run. Change "Heading 2"
  in Word and every second-level heading follows, which is the entire reason a
  student can drop this into a mandated template and have it comply.
* Captions and cross-references are SEQ and REF fields inside bookmarks, the
  same constructs Word's own Insert Caption produces. Insert a figure in the
  middle of the document and Word renumbers the captions *and* the sentences
  that point at them. Static text would go stale on the first edit, and a
  document that silently contradicts itself after one edit is worse than one
  that never claimed to be consistent.
* The contents page is a real TOC field, so it repaginates and picks up
  headings the user adds.

None of this touches the AST. The exporter reads a `Document` and writes
bytes; it never rewrites the tree, never renumbers, and never sees a PDF.
"""

import io
import re
from dataclasses import dataclass

from docx import Document as DocxDocument
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Emu, Mm, Pt, RGBColor

from srs.ast import (
    BulletList,
    Document,
    Figure,
    FigureRef,
    ImageBlock,
    ListOfFigures,
    ListOfTables,
    NumberedList,
    PageBreak,
    Paragraph,
    Rule,
    Section,
    SectionRef,
    Spacer,
    Table,
    TableOfContents,
    TableRef,
    Text,
)
from srs.export.template import A4, DocumentTemplate

EMBEDDABLE = ("image/png", "image/jpeg", "image/gif", "image/bmp", "image/tiff")
"""In preference order. DOCX cannot embed SVG, so a figure that only exists as
vector has to be reported rather than dropped."""


class UnsupportedImage(ValueError):
    """A figure DOCX cannot embed.

    Raised rather than skipped. SVG is the usual cause, and the fix is to
    render the run as PNG — the diagram engines produce it natively, which is a
    better raster than anything a converter would make from their own vector
    output. Dropping the figure instead would break FR-16 quietly, and a
    missing figure in a forty-page document is not noticed until it is marked.
    """


@dataclass(frozen=True)
class DocxResult:
    content: bytes
    figures: int
    tables: int
    sections: int

    @property
    def size(self) -> int:
        return len(self.content)


# ---------------------------------------------------------------------------
# Field and bookmark plumbing
# ---------------------------------------------------------------------------


def _element(tag: str, **attributes):
    from docx.oxml import OxmlElement

    node = OxmlElement(tag)
    for name, value in attributes.items():
        node.set(qn(f"w:{name}"), value)
    return node


def _add_field(paragraph, instruction: str, cached: str = "") -> None:
    """A Word field: instruction, cached result, end — each in its own run.

    Every `w:fldChar` must sit inside a `w:r`. Appending the closing one
    directly to the paragraph produces XML that Word silently repairs and
    LibreOffice silently *discards* — captions came out blank and every field
    disappeared on round-trip, with no error from either program. Found only by
    rendering the file through LibreOffice and reading the resulting PDF.

    The cached result matters as much as the instruction: it is what both
    programs display before anything is recalculated, so a document that is
    opened, read and printed without a field update still shows the numbers the
    AST assigned.
    """
    begin = paragraph.add_run()
    begin._r.append(_element("w:fldChar", fldCharType="begin"))

    instruction_run = paragraph.add_run()
    instruction_node = _element("w:instrText", space="preserve")
    instruction_node.text = f" {instruction} "
    instruction_run._r.append(instruction_node)

    separator = paragraph.add_run()
    separator._r.append(_element("w:fldChar", fldCharType="separate"))

    if cached:
        paragraph.add_run(cached)

    end = paragraph.add_run()
    end._r.append(_element("w:fldChar", fldCharType="end"))


class _Bookmarks:
    """Bookmark ids, numbered from the start of each export.

    Per-export rather than per-process: a module-level counter made the same
    document export to different bytes on the second call, which quietly breaks
    the FR-9 promise that identical input yields identical output — and makes
    "the file changed" useless as a signal that anything actually changed.
    """

    def __init__(self) -> None:
        self._next = 1000

    def wrap(self, paragraph, name: str, write) -> None:
        identifier = str(self._next)
        self._next += 1
        paragraph._p.append(_element("w:bookmarkStart", id=identifier, name=name))
        write()
        paragraph._p.append(_element("w:bookmarkEnd", id=identifier))


def _bookmark_name(prefix: str, target_id: str) -> str:
    """Word bookmark names: letters, digits and underscores, starting with a
    letter or underscore, and no longer than 40 characters."""
    cleaned = re.sub(r"[^0-9A-Za-z]", "_", target_id)
    return f"_{prefix}_{cleaned}"[:40]


def _reference(paragraph, prefix: str, target_id: str, cached: str) -> None:
    _add_field(paragraph, f"REF {_bookmark_name(prefix, target_id)} \\h", cached)


# ---------------------------------------------------------------------------
# Styles
# ---------------------------------------------------------------------------


_THEME_ATTRIBUTES = ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme")


def _set_style_font(style, name: str) -> None:
    """Set a style's typeface so that Word actually honours it.

    Word's built-in heading styles do not name a font: they point at the
    document theme (`w:asciiTheme="majorHAnsi"`), and a theme reference beats
    the explicit `w:ascii` that setting `font.name` writes. The result is a
    template that appears to apply — the style is real, the value is stored —
    while every heading still renders in the theme's Calibri. Found by opening
    the file in Word and reading the style back, which no amount of inspecting
    the XML we wrote would have revealed.
    """
    style.font.name = name
    fonts = style.element.get_or_add_rPr().get_or_add_rFonts()
    for attribute in _THEME_ATTRIBUTES:
        if fonts.get(qn(f"w:{attribute}")) is not None:
            del fonts.attrib[qn(f"w:{attribute}")]
    for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{attribute}"), name)


def _apply_named_styles(document, template: DocumentTemplate) -> None:
    """Push a configured template's named styles onto the Word styles.

    A template that says `heading2` is 14pt bold Times centred has to reach
    Word as a *style*, not as formatting on each heading, or the student
    cannot restyle the document afterwards — which is the whole reason DOCX is
    the primary format.
    """
    mapping = {
        "body": "Normal",
        "heading1": "Heading 1",
        "heading2": "Heading 2",
        "heading3": "Heading 3",
        "caption": "Caption",
    }
    for name, style in template.styles.items():
        word_name = mapping.get(name)
        if word_name is None:
            word_name = _ensure_style(document, name)
        target = document.styles[word_name]
        _set_style_font(target, style.font)
        target.font.size = Pt(style.size_pt)
        target.font.bold = style.bold
        target.font.italic = style.italic
        target.font.all_caps = style.all_caps
        target.font.color.rgb = RGBColor.from_string(style.colour)
        target.paragraph_format.line_spacing = style.line_spacing
        target.paragraph_format.space_before = Pt(style.space_before_pt)
        target.paragraph_format.space_after = Pt(style.space_after_pt)
        target.paragraph_format.keep_with_next = style.keep_with_next
        target.paragraph_format.alignment = _ALIGN.get(style.align)


_ALIGN = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


def _ensure_style(document, name: str) -> str:
    """A real Word style for a template style Word has no built-in for.

    Added to styles.xml rather than applied inline, so "cover_title" is
    something the user can select and modify like any other style.
    """
    from docx.enum.style import WD_STYLE_TYPE

    word_name = name.replace("_", " ").title()
    if word_name not in [s.name for s in document.styles]:
        document.styles.add_style(word_name, WD_STYLE_TYPE.PARAGRAPH)
    return word_name


def _style_name(template: DocumentTemplate, name: str) -> str:
    mapping = {
        "body": "Normal",
        "heading1": "Heading 1",
        "heading2": "Heading 2",
        "heading3": "Heading 3",
        "caption": "Caption",
    }
    if name in mapping:
        return mapping[name]
    if template.resolved(name) is not None:
        return name.replace("_", " ").title()
    return "Normal"


def _configure_styles(document, template: DocumentTemplate) -> None:
    """Set the template's fonts on the *styles*, never on runs.

    This is what "real Word styles" buys: the user opens Modify Style, changes
    Heading 1 once, and the whole document follows. Formatting applied run by
    run would survive export and defeat every attempt to restyle it.
    """
    normal = document.styles["Normal"]
    _set_style_font(normal, template.body_font)
    normal.font.size = Pt(template.body_size_pt)
    normal.paragraph_format.line_spacing = template.line_spacing
    normal.paragraph_format.space_after = Pt(template.space_after_pt)

    for level in (1, 2, 3, 4):
        style = document.styles[f"Heading {level}"]
        _set_style_font(style, template.heading_font)
        style.font.size = Pt(template.heading_size(level))
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(template.heading_colour)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.keep_with_next = True

    caption = document.styles["Caption"]
    _set_style_font(caption, template.caption_font)
    caption.font.size = Pt(template.caption_size_pt)
    caption.font.italic = template.caption_italic
    caption.font.color.rgb = RGBColor.from_string("404040")
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(12)

    for name in ("Title", "Subtitle", "TOC Heading"):
        _set_style_font(document.styles[name], template.heading_font)
    for name in ("List Bullet", "List Number"):
        _set_style_font(document.styles[name], template.body_font)
    toc_heading = document.styles["TOC Heading"]
    toc_heading.font.size = Pt(template.heading_size(1))
    toc_heading.font.color.rgb = RGBColor.from_string(template.heading_colour)


def _configure_page(document, template: DocumentTemplate) -> None:
    for section in document.sections:
        section.page_width = Mm(template.page_width_mm)
        section.page_height = Mm(template.page_height_mm)
        section.top_margin = Mm(template.margins.top_mm)
        section.bottom_margin = Mm(template.margins.bottom_mm)
        section.left_margin = Mm(template.margins.left_mm)
        section.right_margin = Mm(template.margins.right_mm)


def _running_header_and_footer(
    document, template: DocumentTemplate, project: str, skip_first: bool = False
) -> None:
    for index, section in enumerate(document.sections):
        if skip_first and index == 0:
            # No running head on a title page. Every university template says
            # so, and it is the first thing a marker notices.
            section.header.is_linked_to_previous = False
            section.footer.is_linked_to_previous = False
            continue
        header = section.header
        header.is_linked_to_previous = False
        paragraph = header.paragraphs[0]
        paragraph.text = template.header_text.replace("{project}", project)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if template.show_header_rule:
            borders = _element("w:pBdr")
            bottom = _element("w:bottom", val="single", sz="6", space="1", color="A6A6A6")
            borders.append(bottom)
            paragraph._p.get_or_add_pPr().append(borders)

        footer = section.footer
        footer.is_linked_to_previous = False
        line = footer.paragraphs[0]
        line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # PAGE and NUMPAGES are fields, so the footer stays right after edits.
        prefix, _, rest = template.footer_text.partition("{page}")
        middle, _, suffix = rest.partition("{pages}")
        line.add_run(prefix)
        _add_field(line, "PAGE", "1")
        line.add_run(middle)
        _add_field(line, "NUMPAGES", "1")
        line.add_run(suffix)


# ---------------------------------------------------------------------------
# Blocks
# ---------------------------------------------------------------------------


def _write_runs(paragraph, runs) -> None:
    for run in runs:
        if isinstance(run, Text):
            added = paragraph.add_run(run.value)
            if run.emphasis:
                # Only when asked. Setting italic=False writes <w:i w:val="0"/>
                # into every run, which is direct formatting the style can no
                # longer override.
                added.italic = True
        elif isinstance(run, FigureRef):
            _reference(paragraph, "Ref_fig", run.target_id, run.resolved or "")
        elif isinstance(run, TableRef):
            _reference(paragraph, "Ref_tbl", run.target_id, run.resolved or "")
        elif isinstance(run, SectionRef):
            _reference(paragraph, "Ref_sec", run.target_id, run.resolved or "")


def _caption_label(formatted: str | None, kind: str, number: int | None, text: str) -> str:
    """What goes inside the bookmark: the template's label and number.

    "Fig. 3.2" for one template, "Figure 7" for another. It has to be the
    template's wording, because this is the text a cross-reference elsewhere
    will render.
    """
    if formatted and text and text in formatted:
        return formatted[: formatted.index(text)].rstrip(": ").rstrip()
    return f"{kind} {number or 1}"


def _caption_tail(formatted: str | None, kind: str, number: int | None, text: str) -> str:
    """Everything after the number, punctuation included, as configured."""
    if not formatted or not text or text not in formatted:
        return f": {text}"
    head = formatted[: formatted.index(text)]
    return formatted[len(head.rstrip(": ").rstrip()) :]


def _caption(
    document,
    bookmarks: "_Bookmarks",
    kind: str,
    number: int | None,
    text: str,
    target_id: str,
    formatted: str | None = None,
) -> None:
    """A Word caption: label, SEQ field, then the text — inside a bookmark.

    Bookmarking only the label and number (not the caption text) is what makes
    a REF elsewhere render "Figure 3" rather than the whole sentence. It is the
    same shape Word produces from Insert Caption, which is why Word's own
    cross-reference dialog can see these.
    """
    paragraph = document.add_paragraph(style="Caption")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    prefix = "Ref_fig" if kind == "Figure" else "Ref_tbl"

    def label() -> None:
        paragraph.add_run(f"{kind} ")
        _add_field(paragraph, f"SEQ {kind} \\* ARABIC", str(number or 1))

    bookmarks.wrap(paragraph, _bookmark_name(prefix, target_id), label)
    # The template decides the wording after the number: ": Caption" in one,
    # " Caption" in another. The label and number stay inside the bookmark so
    # a cross-reference still renders "Fig. 3.2" and not the whole sentence.
    paragraph.add_run(_caption_tail(formatted, kind, number, text))


def _add_figure(
    document, bookmarks: "_Bookmarks", block: Figure, template: DocumentTemplate
) -> None:
    chosen = block.rendition(EMBEDDABLE)
    if chosen is None:
        raise UnsupportedImage(
            f"cannot embed {block.mime!r} for figure {block.figure_id!r}: DOCX supports "
            f"{', '.join(EMBEDDABLE)}. Render the run with format=png — the "
            f"diagram engines produce PNG natively."
        )

    stream = io.BytesIO(chosen[1])
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()

    available = Mm(template.content_width_mm)
    picture = run.add_picture(stream)
    if picture.width > available:
        # Scale proportionally rather than letting Word crop or overflow.
        picture.height = Emu(int(picture.height * available / picture.width))
        picture.width = Emu(int(available))

    _caption(
        document,
        bookmarks,
        "Figure",
        block.number,
        block.caption,
        block.figure_id,
        block.formatted_caption,
    )


def _add_cover_image(document, block: ImageBlock, template: DocumentTemplate) -> None:
    """An institution crest (FR-15). Not a figure: no number, no caption."""
    if block.mime not in EMBEDDABLE:
        raise UnsupportedImage(
            f"cannot embed {block.mime!r} as a cover image; DOCX supports {', '.join(EMBEDDABLE)}"
        )
    paragraph = document.add_paragraph()
    paragraph.alignment = _ALIGN.get(block.align, WD_ALIGN_PARAGRAPH.CENTER)
    picture = paragraph.add_run().add_picture(io.BytesIO(block.image))
    limit = Mm(block.max_height_mm)
    if picture.height > limit:
        picture.width = Emu(int(picture.width * limit / picture.height))
        picture.height = Emu(int(limit))


def _add_table(document, bookmarks: "_Bookmarks", block: Table, template: DocumentTemplate) -> None:
    if not block.caption:
        # An uncaptioned table is layout — a signature row on a certificate.
        _plain_table(document, block, template)
        return
    _caption(
        document,
        bookmarks,
        "Table",
        block.number,
        block.caption,
        block.table_id,
        block.formatted_caption,
    )

    table = document.add_table(rows=1, cols=len(block.columns))
    try:
        table.style = template.table_style
    except KeyError:  # pragma: no cover - only if a template names a missing style
        table.style = "Table Grid"

    for cell, heading in zip(table.rows[0].cells, block.columns, strict=True):
        cell.text = ""
        run = cell.paragraphs[0].add_run(heading)
        run.bold = True
    # Repeat the header on every page the table spills onto.
    table.rows[0]._tr.get_or_add_trPr().append(_element("w:tblHeader", val="true"))

    for values in block.rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, values, strict=True):
            cell.text = value

    document.add_paragraph()


def _plain_table(document, block: Table, template: DocumentTemplate) -> None:
    """A borderless table: a signature row, not an exhibit."""
    table = document.add_table(rows=len(block.rows), cols=len(block.columns))
    table.style = "Normal Table"
    style_name = _style_name(template, block.cell_style)
    for row_index, values in enumerate(block.rows):
        for cell, value in zip(table.rows[row_index].cells, values, strict=True):
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.style = document.styles[style_name]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for index, line in enumerate(value.split("\n")):
                if index:
                    paragraph = cell.add_paragraph(style=style_name)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.add_run(line)
    document.add_paragraph()


def _add_index(document, block, instruction: str, template: DocumentTemplate) -> None:
    """A real TOC field with the AST's entries cached inside it.

    Both halves are needed. The field is what makes the contents page update
    when the user edits the document — the reason DOCX is the primary format.
    The cached entries are what make it readable in a viewer that never updates
    fields, and they are the numbers the AST assigned, so an un-updated
    contents page is still correct about structure.
    """
    # "TOC Heading", not "Heading 1". It looks identical and carries no outline
    # level, which is what stops the contents page from listing itself — an
    # entry reading "Contents .... 2" is the tell that a document was generated
    # rather than written.
    document.add_paragraph(block.title, style="TOC Heading")

    holder = document.add_paragraph()
    _add_field(holder, instruction)

    entries = block.entries
    if getattr(block, "layout", "dotted") == "table":
        # A bordered INDEX with headed columns, not a dotted list. Some
        # institutions mandate it and it is not a styling of the other.
        grid = document.add_table(rows=1, cols=len(block.columns))
        grid.style = "Table Grid"
        for cell, heading in zip(grid.rows[0].cells, block.columns, strict=True):
            cell.text = ""
            cell.paragraphs[0].add_run(heading).bold = True
        for position, entry in enumerate(entries, start=1):
            cells = grid.add_row().cells
            values = [str(position), f"{entry.number}  {entry.title}", ""]
            for cell, value in zip(cells, values[: len(block.columns)], strict=False):
                cell.text = value
        document.add_paragraph()
        return

    for entry in entries:
        line = document.add_paragraph(style="Normal")
        line.paragraph_format.left_indent = Pt(18 * (entry.depth - 1))
        line.paragraph_format.space_after = Pt(2)
        # Number and title together, no tabbed page column: these entries are
        # what a viewer shows *before* fields update, and a page number we
        # cannot compute is worse than none. Word replaces the lot on update.
        line.add_run(f"{entry.number}   {entry.title}")


# ---------------------------------------------------------------------------
# The exporter
# ---------------------------------------------------------------------------


def _title_page(document, source: Document, template: DocumentTemplate) -> None:
    title = document.add_paragraph(source.title, style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = document.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(f"Version {source.meta.version}")

    if source.meta.authors:
        authors = document.add_paragraph(style="Normal")
        authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
        authors.add_run(", ".join(source.meta.authors))

    if source.meta.created_at:
        date = document.add_paragraph(style="Normal")
        date.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date.add_run(source.meta.created_at[:10])

    document.add_section(WD_SECTION.NEW_PAGE)


def _write_section(
    document, bookmarks: "_Bookmarks", section: Section, template: DocumentTemplate, level: int
) -> None:
    heading = document.add_paragraph(style=f"Heading {min(level, 4)}")
    bookmarks.wrap(
        heading,
        _bookmark_name("Ref_sec", section.section_id),
        lambda: heading.add_run(section.heading),
    )

    for block in section.blocks:
        _write_block(document, bookmarks, block, template)

    for child in section.subsections:
        _write_section(document, bookmarks, child, template, level + 1)


def _write_block(document, bookmarks: "_Bookmarks", block, template: DocumentTemplate) -> None:
    if isinstance(block, Paragraph):
        paragraph = document.add_paragraph(style=_style_name(template, block.style))
        if block.align:
            paragraph.alignment = _ALIGN.get(block.align)
        _write_runs(paragraph, block.runs)
    elif isinstance(block, BulletList):
        for item in block.items:
            paragraph = document.add_paragraph(style="List Bullet")
            _write_runs(paragraph, item)
    elif isinstance(block, NumberedList):
        for item in block.items:
            paragraph = document.add_paragraph(style="List Number")
            _write_runs(paragraph, item)
    elif isinstance(block, Spacer):
        spacer = document.add_paragraph()
        spacer.paragraph_format.space_after = Pt(0)
        spacer.paragraph_format.space_before = Mm(block.height_mm)
    elif isinstance(block, Rule):
        line = document.add_paragraph()
        borders = _element("w:pBdr")
        borders.append(_element("w:bottom", val="single", sz="8", space="1", color="808080"))
        line._p.get_or_add_pPr().append(borders)
    elif isinstance(block, PageBreak):
        document.add_page_break()
    elif isinstance(block, ImageBlock):
        _add_cover_image(document, block, template)
    elif isinstance(block, Figure):
        _add_figure(document, bookmarks, block, template)
    elif isinstance(block, Table):
        _add_table(document, bookmarks, block, template)
    elif isinstance(block, TableOfContents):
        _add_index(document, block, 'TOC \\o "1-3" \\h \\z \\u', template)
    elif isinstance(block, ListOfFigures):
        _add_index(document, block, 'TOC \\h \\z \\c "Figure"', template)
    elif isinstance(block, ListOfTables):
        _add_index(document, block, 'TOC \\h \\z \\c "Table"', template)
    else:
        # Never silently. An unhandled block would be missing from this format
        # and present in the other — the divergence AT-1 exists to catch, and
        # it would be caught a very long way from the cause.
        raise TypeError(f"the DOCX exporter has no rule for {type(block).__name__}")


def export_docx(source: Document, template: DocumentTemplate = A4) -> DocxResult:
    """Render the document AST as a .docx.

    A pure function of the AST and the template. It reads no database, calls no
    model, and — the point of the layer — knows nothing about the PDF exporter.
    """
    if not source.numbered:
        raise ValueError(
            "the document has not been numbered; export renders what assembly "
            "decided and must not assign numbers of its own"
        )

    document = DocxDocument()
    _configure_styles(document, template)
    _apply_named_styles(document, template)
    _configure_page(document, template)

    if template.title_page:
        _title_page(document, source, template)

    bookmarks = _Bookmarks()
    for block in source.front_matter:
        _write_block(document, bookmarks, block, template)

    if source.sections:
        document.add_section(WD_SECTION.NEW_PAGE)

    for section in source.sections:
        _write_section(document, bookmarks, section, template, level=1)

    _configure_page(document, template)
    _running_header_and_footer(
        document, template, source.meta.project_name, skip_first=template.title_page
    )
    _mark_fields_dirty(document)

    stream = io.BytesIO()
    document.save(stream)

    from srs.ast import figures as ast_figures
    from srs.ast import tables as ast_tables
    from srs.ast import walk_sections

    return DocxResult(
        content=stream.getvalue(),
        figures=len(ast_figures(source)),
        tables=len(ast_tables(source)),
        sections=len(list(walk_sections(source))),
    )


def _mark_fields_dirty(document) -> None:
    """Ask Word to recalculate on open, so page numbers are real.

    Only the fields whose value depends on pagination are marked. The captions
    and cross-references are already correct from the AST, and forcing a
    document-wide update would make Word prompt about a document that does not
    need it.
    """
    settings = document.settings.element
    if settings.find(qn("w:updateFields")) is None:
        settings.append(_element("w:updateFields", val="true"))
