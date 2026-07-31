"""The DOCX exporter, checked at the level of the file it writes.

These assertions read the OOXML back out of the produced .docx, because the
things that matter here are things python-docx will happily let you get wrong
and neither Word nor LibreOffice will complain about. Two of the bugs pinned
below were found by opening the output in Word and LibreOffice and were
invisible from the Python side:

* setting a font on a built-in heading style has no effect, because the style
  points at the document theme and the theme reference wins;
* a `w:fldChar` appended to a paragraph instead of a run makes every field
  vanish in LibreOffice — blank captions, no page numbers — while Word repairs
  it silently and looks fine.
"""

import re
import zipfile
from io import BytesIO

import pytest
from conftest import sample_png
from docx import Document as DocxDocument

from cpm.fixtures import load_library_management_system
from srs.assemble import assemble_srs
from srs.ast import figures as ast_figures
from srs.export.docx import EMBEDDABLE, UnsupportedImage, export_docx
from srs.export.template import A4, IEEE_TIMES, US_LETTER, DocumentTemplate, get_template
from srs.ieee830 import CAPTIONS, FigureInput

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

PNG = sample_png()


def inputs(*types: str) -> list[FigureInput]:
    return [
        FigureInput(diagram_type=t, title=CAPTIONS.get(t, t), image=PNG, mime="image/png")
        for t in (types or ("class", "use_case"))
    ]


@pytest.fixture
def cpm():
    return load_library_management_system()


@pytest.fixture
async def document(cpm):
    return (await assemble_srs(cpm, inputs())).document


@pytest.fixture
def exported(document):
    return export_docx(document, A4)


def xml_of(result, part: str = "word/document.xml") -> str:
    with zipfile.ZipFile(BytesIO(result.content)) as archive:
        return archive.read(part).decode("utf-8")


def parts_of(result) -> list[str]:
    with zipfile.ZipFile(BytesIO(result.content)) as archive:
        return archive.namelist()


def reopened(result):
    return DocxDocument(BytesIO(result.content))


# --------------------------------------------------------------------------
# It is a real .docx
# --------------------------------------------------------------------------


def test_the_output_is_a_readable_package(exported) -> None:
    parts = parts_of(exported)
    for required in (
        "[Content_Types].xml",
        "word/document.xml",
        "word/styles.xml",
        "word/settings.xml",
    ):
        assert required in parts

    document = reopened(exported)
    assert document.paragraphs
    assert len(document.tables) == 15


def test_every_relationship_target_exists(exported) -> None:
    # A dangling relationship is what makes Word show its repair prompt.
    import posixpath

    parts = set(parts_of(exported))
    with zipfile.ZipFile(BytesIO(exported.content)) as archive:
        checked = 0
        for name in [n for n in parts if n.endswith(".rels")]:
            base = name.rsplit("_rels/", 1)[0]
            body = archive.read(name).decode("utf-8")
            for target, mode in re.findall(r'Target="([^"]+)"(?:\s+TargetMode="([^"]+)")?', body):
                if mode == "External" or target.startswith("http"):
                    continue
                resolved = posixpath.normpath(posixpath.join(base, target))
                assert resolved in parts, f"{name} points at missing {resolved}"
                checked += 1
        assert checked > 5, "the relationship scan found almost nothing to check"


# --------------------------------------------------------------------------
# Real styles
# --------------------------------------------------------------------------


def test_headings_use_word_heading_styles(exported) -> None:
    document = reopened(exported)
    used = {p.style.name for p in document.paragraphs}
    assert {"Heading 1", "Heading 2", "Heading 3", "Caption", "Title"} <= used


def test_the_outline_carries_the_ast_numbers(exported, document) -> None:
    headings = [p.text for p in reopened(exported).paragraphs if p.style.name.startswith("Heading")]
    assert "1 Introduction" in headings
    assert "1.1 Purpose" in headings
    assert "3.2.1 Borrow Book" in headings


def test_nothing_is_formatted_by_hand(exported) -> None:
    # If body text carried direct formatting, restyling in Word would not work
    # — which is the entire reason to use styles rather than imitate them.
    document = reopened(exported)
    body = [p for p in document.paragraphs if p.style.name == "Normal"]
    assert body
    for paragraph in body:
        for run in paragraph.runs:
            assert run.font.size is None, paragraph.text[:60]
            assert run.font.name is None, paragraph.text[:60]


def test_the_template_font_reaches_the_style_and_not_only_the_api(exported) -> None:
    # The bug this pins: python-docx stores font.name happily while Word keeps
    # rendering the theme font, because w:asciiTheme outranks w:ascii.
    styles = xml_of(exported, "word/styles.xml")
    heading = re.search(r'<w:style [^>]*w:styleId="Heading1".*?</w:style>', styles, re.S)
    assert heading, "Heading 1 is not defined in the package"
    assert "asciiTheme" not in heading.group(0), "the theme reference still overrides the template"
    assert f'w:ascii="{A4.heading_font}"' in heading.group(0)


def test_a_different_template_produces_a_different_document(document) -> None:
    a4 = xml_of(export_docx(document, A4), "word/styles.xml")
    times = xml_of(export_docx(document, IEEE_TIMES), "word/styles.xml")

    assert 'w:ascii="Times New Roman"' in times
    assert 'w:ascii="Times New Roman"' not in a4
    # 12pt body is stored as half-points.
    assert 'w:val="24"' in times


def test_the_contents_heading_does_not_list_itself(exported) -> None:
    document = reopened(exported)
    styles = {
        p.text: p.style.name
        for p in document.paragraphs
        if p.text in ("Contents", "List of Figures", "List of Tables")
    }
    assert styles == {
        "Contents": "TOC Heading",
        "List of Figures": "TOC Heading",
        "List of Tables": "TOC Heading",
    }


# --------------------------------------------------------------------------
# Page setup from the template
# --------------------------------------------------------------------------


def test_page_size_and_margins_come_from_the_template(document) -> None:
    for template, width_mm in ((A4, 210.0), (US_LETTER, 215.9)):
        section = reopened(export_docx(document, template)).sections[0]
        assert round(section.page_width.mm, 1) == width_mm
        assert round(section.top_margin.mm, 1) == round(template.margins.top_mm, 1)
        assert round(section.left_margin.mm, 1) == round(template.margins.left_mm, 1)


def test_an_unknown_template_setting_is_refused_rather_than_ignored() -> None:
    with pytest.raises(ValueError, match="hedaing_font"):
        DocumentTemplate.from_dict({"hedaing_font": "Arial"})


def test_a_template_can_be_built_from_stored_configuration() -> None:
    template = DocumentTemplate.from_dict(
        {
            "name": "Course template",
            "body_font": "Arial",
            "body_size_pt": 12,
            "heading_sizes_pt": [18, 15, 13],
            "margins": {"left_mm": 38.1, "right_mm": 25.4},
        }
    )
    assert template.body_font == "Arial"
    assert template.heading_size(2) == 15
    assert template.margins.left_mm == 38.1
    assert get_template("ieee-times").body_font == "Times New Roman"


# --------------------------------------------------------------------------
# Figures and captions
# --------------------------------------------------------------------------


def test_every_figure_is_embedded_as_a_drawing(exported, document) -> None:
    # Counted as drawings rather than media parts: OOXML stores one part per
    # distinct image and reuses it, which is correct behaviour and would make a
    # part count lie whenever two figures happen to be byte-identical.
    assert xml_of(exported).count("<w:drawing>") == len(ast_figures(document))
    assert exported.figures == len(ast_figures(document))
    assert [p for p in parts_of(exported) if p.startswith("word/media/")]


def test_each_figure_has_a_caption_bound_to_it(exported, document) -> None:
    captions = [p for p in reopened(exported).paragraphs if p.style.name == "Caption"]
    figure_captions = [p.text for p in captions if p.text.startswith("Figure")]

    assert len(figure_captions) == len(ast_figures(document))
    for figure in ast_figures(document):
        assert any(
            text.startswith(f"Figure {figure.number}: ") and figure.caption in text
            for text in figure_captions
        ), figure.caption


def test_captions_are_seq_fields_so_word_renumbers_them(exported) -> None:
    body = xml_of(exported)
    assert body.count("SEQ Figure") == exported.figures
    assert body.count("SEQ Table") == exported.tables


def test_every_field_char_lives_inside_a_run(exported) -> None:
    # The LibreOffice bug, pinned. A fldChar directly under w:p makes every
    # field disappear silently in LibreOffice while Word repairs it.
    import xml.etree.ElementTree as ET

    with zipfile.ZipFile(BytesIO(exported.content)) as archive:
        for part in [n for n in archive.namelist() if n.endswith(".xml") and "word/" in n]:
            root = ET.fromstring(archive.read(part))
            for parent in root.iter():
                if parent.tag == f"{{{W}}}r":
                    continue
                for child in parent:
                    assert child.tag != f"{{{W}}}fldChar", (
                        f"{part}: a fldChar sits directly in <{parent.tag.split('}')[-1]}>"
                    )


def test_a_cross_reference_is_a_ref_field_pointing_at_a_real_bookmark(exported) -> None:
    body = xml_of(exported)
    referenced = set(re.findall(r"REF (_Ref_\w+)", body))
    declared = set(re.findall(r'w:bookmarkStart[^>]*w:name="([^"]+)"', body))

    assert referenced, "no cross-references were emitted at all"
    assert referenced <= declared, referenced - declared


def test_the_cached_reference_text_is_the_number_the_ast_assigned(exported, document) -> None:
    body = reopened(exported)
    text = " ".join(p.text for p in body.paragraphs)
    use_case = next(f for f in ast_figures(document) if f.diagram_type == "use_case")
    assert f"shown in Figure {use_case.number}" in text


def test_svg_is_refused_with_the_fix_named(cpm) -> None:
    async def build():
        return (
            await assemble_srs(cpm, [FigureInput("class", "Class", b"<svg/>", "image/svg+xml")])
        ).document

    import asyncio

    document = asyncio.get_event_loop_policy().new_event_loop().run_until_complete(build())
    with pytest.raises(UnsupportedImage) as excinfo:
        export_docx(document, A4)

    assert "format=png" in str(excinfo.value)
    assert "image/svg+xml" not in EMBEDDABLE


# --------------------------------------------------------------------------
# Headers, footers, contents
# --------------------------------------------------------------------------


def test_running_header_and_footer_are_present_on_body_sections(exported) -> None:
    document = reopened(exported)
    assert len(document.sections) >= 2

    body_sections = document.sections[1:]
    for section in body_sections:
        assert "Library Management System" in section.header.paragraphs[0].text
        footer = section.footer.paragraphs[0]
        assert "Page" in footer.text


def test_the_title_page_carries_no_running_head(exported) -> None:
    first = reopened(exported).sections[0]
    assert first.header.paragraphs[0].text.strip() == ""
    assert first.footer.paragraphs[0].text.strip() == ""


def test_the_footer_page_number_is_a_field_not_a_literal(exported) -> None:
    footers = [p for p in parts_of(exported) if re.match(r"word/footer\d+\.xml", p)]
    combined = "".join(xml_of(exported, part) for part in footers)
    assert "PAGE" in combined and "NUMPAGES" in combined


def test_the_contents_is_a_toc_field_with_the_entries_cached(exported, document) -> None:
    body = xml_of(exported)
    assert 'TOC \\o "1-3"' in body
    assert 'TOC \\h \\z \\c "Figure"' in body
    assert 'TOC \\h \\z \\c "Table"' in body

    # Cached entries so a viewer that never updates fields still shows them.
    text = " ".join(p.text for p in reopened(exported).paragraphs)
    for entry in document.front_matter[0].entries[:5]:
        assert f"{entry.number}   {entry.title}" in text


def test_fields_are_marked_for_update_on_open(exported) -> None:
    assert "updateFields" in xml_of(exported, "word/settings.xml")


# --------------------------------------------------------------------------
# The layer boundary
# --------------------------------------------------------------------------


def test_the_exporter_refuses_an_unnumbered_document(cpm) -> None:
    import asyncio

    from srs.ieee830 import build_document
    from srs.prose import DeterministicProse

    loop = asyncio.get_event_loop_policy().new_event_loop()
    raw = loop.run_until_complete(build_document(cpm, inputs(), DeterministicProse()))

    with pytest.raises(ValueError, match="not been numbered"):
        export_docx(raw, A4)


def test_the_docx_exporter_knows_nothing_about_pdf() -> None:
    import ast
    from pathlib import Path

    import srs.export.docx as module

    tree = ast.parse(Path(module.__file__).read_text(encoding="utf-8"))
    imported = set()
    for node in ast.walk(tree):
        if isinstance(node, ast.Import):
            imported.update(alias.name.split(".")[0] for alias in node.names)
        elif isinstance(node, ast.ImportFrom):
            imported.add((node.module or "").split(".")[0])

    assert not imported & {"reportlab", "weasyprint", "fpdf", "pdfkit", "fitz"}
    assert "srs" in imported, "the exporter should read the AST and nothing else"


def test_export_is_deterministic(document) -> None:
    # Same tree, same template, same file — bookmark ids included.
    first = export_docx(document, A4)
    second = export_docx(document, A4)
    assert xml_of(first) == xml_of(second)
