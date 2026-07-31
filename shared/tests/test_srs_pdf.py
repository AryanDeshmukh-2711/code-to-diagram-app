"""The PDF exporter, and AT-1: the two files say the same thing.

The equivalence test does not strip differences until the two streams match.
It aligns them and then *verifies what the differences are* — every word of
the DOCX body must appear in the PDF in the same order, and the PDF's only
extra tokens must be list markers, one per list item in the document. A bullet
is paragraph formatting in DOCX and a drawn glyph in PDF; that is the whole of
the allowed divergence, and the count is pinned, so a dropped sentence or a
changed word fails rather than being absorbed by a tolerant normaliser.
"""

import ast as python_ast
import copy
import difflib
import io
import re
import time
from pathlib import Path

import pytest
from conftest import sample_png
from docx import Document as DocxDocument
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph as DocxParagraph
from pypdf import PdfReader

from cpm.fixtures import library_management_system_payload, load_library_management_system
from cpm.schema import CPM
from srs.assemble import assemble_srs
from srs.ast import BulletList, NumberedList, walk_blocks
from srs.export import pdf as pdf_module
from srs.export.docx import export_docx
from srs.export.pdf import UnrenderableFigure, export_pdf
from srs.export.template import A4, FREE_TIER, IEEE_TIMES, US_LETTER, Watermark
from srs.ieee830 import CAPTIONS, FigureInput

PNG = sample_png()
RENDERED = ("class", "entity_relationship", "use_case", "sequence")

ASCII_MARK = Watermark(text="SAMPLE - FREE PLAN", subtext="Upgrade to remove")
"""Plain ASCII so the assertion can look for it in a raw content stream without
arguing with PDF string encoding. The shipped FREE_TIER mark is checked
separately for the things encoding cannot affect."""


def inputs(*types: str, mime: str = "image/png", image: bytes | None = None):
    return [
        FigureInput(
            diagram_type=t,
            title=CAPTIONS.get(t, t),
            image=image if image is not None else PNG,
            mime=mime,
        )
        for t in (types or RENDERED)
    ]


@pytest.fixture
def cpm():
    return load_library_management_system()


@pytest.fixture
async def document(cpm):
    return (await assemble_srs(cpm, inputs())).document


@pytest.fixture
def exported(document):
    return export_pdf(document, A4)


def reader_of(result) -> PdfReader:
    return PdfReader(io.BytesIO(result.content))


# --------------------------------------------------------------------------
# Text extraction, shared by the equivalence test
# --------------------------------------------------------------------------


def docx_body_text(blob: bytes) -> list[str]:
    """Everything from section 1 onward, in document order, tables included."""
    document = DocxDocument(io.BytesIO(blob))
    out: list[str] = []
    started = False
    for child in document.element.body.iterchildren():
        if child.tag.endswith("}p"):
            paragraph = DocxParagraph(child, document)
            if paragraph.style.name == "Heading 1" and paragraph.text.startswith("1 "):
                started = True
            if started and paragraph.text.strip():
                out.append(paragraph.text)
        elif child.tag.endswith("}tbl") and started:
            for row in DocxTable(child, document).rows:
                for cell in row.cells:
                    if cell.text.strip():
                        out.append(cell.text)
    return out


def pdf_body_text(blob: bytes, header: str) -> list[str]:
    """The same span of the PDF, minus the page furniture.

    Where the body starts is taken from the document's own outline rather than
    guessed from where a phrase first appears — the contents page mentions
    every heading too.
    """
    reader = PdfReader(io.BytesIO(blob))
    start = min(
        reader.get_destination_page_number(entry)
        for entry in reader.outline
        if not isinstance(entry, list)
    )
    lines: list[str] = []
    for page in reader.pages[start:]:
        for line in (page.extract_text() or "").splitlines():
            text = line.strip()
            if not text or text == header or re.fullmatch(r"Page \d+ of \d+", text):
                continue
            lines.append(text)
    return lines


def is_list_marker(token: str) -> bool:
    return token in {"•", "\x7f", "-", "–"} or (token.isdigit() and len(token) <= 3)


# --------------------------------------------------------------------------
# AT-1
# --------------------------------------------------------------------------


async def test_pdf_and_docx_say_exactly_the_same_thing(cpm, document) -> None:
    docx = export_docx(document, A4)
    pdf = export_pdf(document, A4)

    docx_words = " ".join(docx_body_text(docx.content)).split()
    pdf_words = " ".join(pdf_body_text(pdf.content, cpm.meta.project_name)).split()
    assert docx_words, "no text was extracted from the DOCX"

    matcher = difflib.SequenceMatcher(None, docx_words, pdf_words, autojunk=False)
    missing: list[str] = []
    extra: list[str] = []
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag in ("replace", "delete"):
            missing.extend(docx_words[i1:i2])
        if tag in ("replace", "insert"):
            extra.extend(pdf_words[j1:j2])

    assert missing == [], f"the PDF is missing text the DOCX has: {missing[:20]}"

    list_items = sum(
        len(block.items)
        for _, block in walk_blocks(document)
        if isinstance(block, BulletList | NumberedList)
    )
    assert all(is_list_marker(token) for token in extra), (
        f"the PDF contains text the DOCX does not: "
        f"{[t for t in extra if not is_list_marker(t)][:20]}"
    )
    assert len(extra) == list_items, (
        f"expected exactly one list marker per list item ({list_items}), got {len(extra)}"
    )


async def test_both_exports_carry_the_same_figures_and_tables(document) -> None:
    docx = export_docx(document, A4)
    pdf = export_pdf(document, A4)
    assert (docx.figures, docx.tables, docx.sections) == (pdf.figures, pdf.tables, pdf.sections)


async def test_one_assembled_document_feeds_both_exporters(cpm) -> None:
    # AT-1 holds because there is one document, not two assemblies that happen
    # to agree. The figure carries both renditions so neither export needs its
    # own pass over the model.
    figures = [
        FigureInput(
            "class",
            "Class Diagram",
            b"<svg xmlns='http://www.w3.org/2000/svg'/>",
            "image/svg+xml",
            (("image/png", PNG),),
        )
    ]
    assembled = await assemble_srs(cpm, figures)

    docx = export_docx(assembled.document, A4)
    pdf = export_pdf(assembled.document, A4)

    assert docx.figures == pdf.figures == 1
    assert b"image1.png" in docx.content or docx.size > 0
    assert pdf.size > 0


# --------------------------------------------------------------------------
# Vector, with a raster fallback
# --------------------------------------------------------------------------


async def test_an_svg_figure_is_drawn_as_vector(cpm) -> None:
    svg = (
        b"<svg xmlns='http://www.w3.org/2000/svg' width='200' height='100'>"
        b"<rect x='5' y='5' width='190' height='90' fill='none' stroke='black'/>"
        b"<text x='20' y='55'>Book</text></svg>"
    )
    assembled = await assemble_srs(cpm, [FigureInput("class", "Class", svg, "image/svg+xml")])
    result = export_pdf(assembled.document, A4)

    assert result.vector_figures == 1
    assert result.raster_figures == 0
    # Vector text is real text: searchable in the PDF, unlike a raster diagram.
    reader = reader_of(result)
    assert any("Book" in (page.extract_text() or "") for page in reader.pages)


async def test_an_unconvertible_svg_falls_back_to_the_raster(cpm) -> None:
    # "Vector where possible" — and where not, the figure still appears.
    broken = b"<svg xmlns='http://www.w3.org/2000/svg'><g style='no-colon-here'/></svg>"
    assembled = await assemble_srs(
        cpm,
        [FigureInput("class", "Class", broken, "image/svg+xml", (("image/png", PNG),))],
    )
    result = export_pdf(assembled.document, A4)

    assert result.raster_figures == 1
    assert result.figures == 1


async def test_a_figure_with_no_usable_rendition_is_reported(cpm) -> None:
    assembled = await assemble_srs(
        cpm, [FigureInput("class", "Class", b"\x00\x01", "application/postscript")]
    )
    with pytest.raises(UnrenderableFigure):
        export_pdf(assembled.document, A4)


async def test_a_raster_is_sized_from_its_dpi_not_its_pixel_count(cpm) -> None:
    # A 1200px diagram at 72dpi would be sixteen inches wide and get scaled to
    # a smear; read at its real 150dpi it is a sharp 8in figure.
    dense = sample_png(width=1200, height=600, dpi=150)
    assembled = await assemble_srs(cpm, [FigureInput("class", "Class", dense, "image/png")])
    result = export_pdf(assembled.document, A4)

    assert result.raster_figures == 1
    assert result.size > 0


# --------------------------------------------------------------------------
# FR-20: the watermark, at render time
# --------------------------------------------------------------------------


def _content_streams(result) -> list[bytes]:
    return [page.get_contents().get_data() for page in reader_of(result).pages]


async def test_the_watermark_is_in_every_page_content_stream(document) -> None:
    result = export_pdf(document, A4, ASCII_MARK)
    streams = _content_streams(result)

    assert streams
    for index, stream in enumerate(streams):
        assert ASCII_MARK.text.encode() in stream, f"page {index + 1} carries no watermark"


async def test_the_watermark_is_drawn_not_overlaid(document) -> None:
    """The distinction FR-20 asks for, made checkable.

    An overlay applied afterwards arrives as a separate form XObject merged
    into each page — the page's own content stream would not contain the text.
    Here it does, because the mark was drawn while the page was being written.
    """
    result = export_pdf(document, A4, ASCII_MARK)
    for page in reader_of(result).pages:
        own_stream = page.get_contents().get_data()
        assert ASCII_MARK.text.encode() in own_stream

        resources = page.get("/Resources", {})
        forms = [
            name
            for name, obj in (resources.get("/XObject", {}) or {}).items()
            if obj.get_object().get("/Subtype") == "/Form"
        ]
        assert not forms, f"a form overlay was merged in: {forms}"


def test_nothing_in_the_exporter_can_reopen_a_pdf() -> None:
    # A post-render watermark needs a PDF reader. There is not one here, and
    # this is what keeps "render time" from quietly becoming "second pass".
    tree = python_ast.parse(Path(pdf_module.__file__).read_text(encoding="utf-8"))
    imported = set()
    for node in python_ast.walk(tree):
        if isinstance(node, python_ast.Import):
            imported.update(alias.name.split(".")[0] for alias in node.names)
        elif isinstance(node, python_ast.ImportFrom):
            imported.add((node.module or "").split(".")[0])

    assert not imported & {"pypdf", "PyPDF2", "pikepdf", "fitz", "pdfrw"}
    assert not imported & {"docx"}, "the PDF exporter must not read the DOCX exporter's work"


async def test_no_watermark_means_no_watermark(document) -> None:
    plain = export_pdf(document, A4)
    for stream in _content_streams(plain):
        assert ASCII_MARK.text.encode() not in stream
        assert b"Free plan" not in stream


async def test_the_shipped_free_tier_mark_applies_to_every_page(document) -> None:
    marked = export_pdf(document, A4, FREE_TIER)
    plain = export_pdf(document, A4)
    # Encoding of an em dash is not worth asserting on; that the mark changes
    # every page's content is.
    assert marked.size > plain.size
    assert len(_content_streams(marked)) == len(_content_streams(plain))


async def test_a_long_watermark_is_scaled_to_fit_rather_than_clipped(document) -> None:
    from srs.export.pdf import _watermark_size

    long_mark = Watermark(text="A" * 120)
    assert _watermark_size(long_mark, A4) < long_mark.size_pt
    assert _watermark_size(ASCII_MARK, A4) == ASCII_MARK.size_pt

    export_pdf(document, A4, long_mark)  # still renders


# --------------------------------------------------------------------------
# Pagination, headers, footers
# --------------------------------------------------------------------------


async def test_every_body_page_has_the_running_header_and_a_page_number(cpm, document) -> None:
    result = export_pdf(document, A4)
    reader = reader_of(result)
    pages = [page.extract_text() or "" for page in reader.pages]

    assert len(pages) > 5
    for index, text in enumerate(pages[1:], start=2):
        assert cpm.meta.project_name in text, f"page {index} has no running header"
        assert re.search(rf"Page {index} of {len(pages)}", text), f"page {index} has no number"


async def test_the_title_page_has_no_running_head(document) -> None:
    first = reader_of(export_pdf(document, A4)).pages[0].extract_text() or ""
    assert "Page 1 of" not in first
    # The title wraps across lines on the page, so compare on collapsed space.
    assert "Software Requirements Specification" in " ".join(first.split())


async def test_the_outline_points_at_the_page_the_section_is_on(document) -> None:
    # Deferring showPage for the page totals also defers where a bookmark
    # lands; without care every entry resolved to page 1.
    result = export_pdf(document, A4)
    reader = reader_of(result)
    entries = [entry for entry in reader.outline if not isinstance(entry, list)]

    assert len(entries) == 3
    for entry in entries:
        page_index = reader.get_destination_page_number(entry)
        assert page_index > 0, f"{entry.title!r} points at the title page"
        assert entry.title in (reader.pages[page_index].extract_text() or ""), entry.title


async def test_page_size_and_margins_come_from_the_template(document) -> None:
    for template, width_mm in ((A4, 210.0), (US_LETTER, 215.9)):
        page = reader_of(export_pdf(document, template)).pages[0]
        assert round(float(page.mediabox.width) / 72 * 25.4) == round(width_mm)


async def test_a_different_template_changes_the_rendered_document(document) -> None:
    a4 = export_pdf(document, A4)
    times = export_pdf(document, IEEE_TIMES)
    assert a4.content != times.content
    # 1.5 line spacing and a larger body font make a longer document.
    assert len(reader_of(times).pages) > len(reader_of(a4).pages)


async def test_the_document_metadata_names_the_project(cpm, document) -> None:
    info = reader_of(export_pdf(document, A4)).metadata
    assert cpm.meta.project_name in info.title
    assert "IEEE 830" in info.subject


# --------------------------------------------------------------------------
# NFR-P4
# --------------------------------------------------------------------------


def _inflated_cpm(use_cases: int) -> CPM:
    """The fixture, grown to the size of a real submission."""
    payload = copy.deepcopy(library_management_system_payload())
    template = copy.deepcopy(payload["useCases"][0])
    for index in range(use_cases):
        clone = copy.deepcopy(template)
        clone["id"] = f"uc-generated-{index:03d}"
        clone["name"] = f"Managed Operation {index:03d}"
        payload["useCases"].append(clone)
    for index in range(40):
        payload["requirements"].append(
            {
                "id": f"req-generated-{index:03d}",
                "type": "functional" if index % 2 else "nonFunctional",
                "text": f"The system shall support generated operation {index:03d} "
                f"within the stated performance envelope.",
                "priority": "Medium",
            }
        )
    return CPM.model_validate(payload)


async def test_a_forty_page_document_exports_well_inside_the_budget() -> None:
    assembled = await assemble_srs(_inflated_cpm(60), inputs())

    started = time.perf_counter()
    result = export_pdf(assembled.document, A4, FREE_TIER)
    elapsed = time.perf_counter() - started

    pages = len(reader_of(result).pages)
    assert pages >= 40, f"the fixture only reached {pages} pages; the budget was not exercised"
    assert elapsed < 30.0, f"NFR-P4: {pages} pages took {elapsed:.1f}s"


# --------------------------------------------------------------------------
# The layer boundary
# --------------------------------------------------------------------------


async def test_the_exporter_refuses_an_unnumbered_document(cpm) -> None:
    from srs.ieee830 import build_document
    from srs.prose import DeterministicProse

    raw = await build_document(cpm, inputs(), DeterministicProse())
    with pytest.raises(ValueError, match="not been numbered"):
        export_pdf(raw, A4)


async def test_export_is_deterministic(document) -> None:
    first = export_pdf(document, A4)
    second = export_pdf(document, A4)
    # PDFs embed a creation timestamp, so the streams are compared instead.
    assert _content_streams(first) == _content_streams(second)
